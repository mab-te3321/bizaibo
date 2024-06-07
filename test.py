from docx import Document

def add_price_to_table(docx_path, row_index, price):
    # Load the existing Word document
    doc = Document(docx_path)
    
    # Assuming the table you want to modify is the first one in the document
    table = doc.tables[0]
    
    # Check if the specified row index is within the table's bounds
# Iterate over paragraphs and print index, text, and style
    for idx, paragraph in enumerate(doc.paragraphs):
        print(f"Paragraph {idx}:")
        print(f"  Text: {paragraph.text}")
        print(f"  Style: {paragraph.style.name}")

    # Iterate over tables and print index, text, and style of each cell
    for idx, table in enumerate(doc.tables):
        print(f"Table {idx}:")
        for row in table.rows:
            for cell in row.cells:
                print(f"  Cell text: {cell.text}")
                print(f"  Cell style: {cell.paragraphs[0].style.name}")

    
    # Save the modified document
    # doc.save(docx_path)

# Usage
docx_file_path = r'media\invoices\template.docx'
row_to_modify = 1  # Index of the row where you want to add/modify the price
price_to_add = 'MAB'  # Price to add
add_price_to_table(docx_file_path, row_to_modify, price_to_add)
