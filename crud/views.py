from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.apps import apps
from django.urls import reverse_lazy
from . import forms
from .forms import ClientFormSet, InvoiceFormSet
from .models import *
from .serializers import *
from rest_framework import serializers
from rest_framework import viewsets
from rest_framework.response import Response
from django.db.models import Q
from django.core.exceptions import ImproperlyConfigured
from django.shortcuts import render,redirect
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
# Create your views here.
from django.http import HttpResponse
from crud.task import add 
import importlib
from rest_framework.exceptions import NotFound
from django.conf import settings
from django.http import FileResponse
import os
from django.http import Http404
# imports for docx
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
class DynamicModelViewSet(viewsets.ModelViewSet):
    def get_queryset(self):
        model_name = self.kwargs['model_name']
        try:
            model = apps.get_model('crud', model_name)
        except LookupError:
            raise NotFound(f"Model '{model_name}' not found.")
        return model.objects.all()

    def get_serializer_class(self):
        model_name = self.kwargs['model_name']
        module_path = 'crud.serializers'
        serializer_class_name = model_name + 'Serializer'

        try:
            serializers_module = importlib.import_module(module_path)
            serializer_class = getattr(serializers_module, serializer_class_name)
        except (ImportError, AttributeError):
            raise NotFound(f"Serializer for model '{model_name}' not found.")
        
        return serializer_class
import tempfile
def index(request):
    file_path = os.path.join(settings.BASE_DIR, 'test.rest')  # Update this to the path of your file

    if os.path.exists(file_path):
        return FileResponse(open(file_path, 'rb'), as_attachment=True, filename='test.rest')  # 'as_attachment=True' makes it a download
    else:
        raise Http404("File not found.")
def manage_items(request):
    if request.method == 'POST':
        formset = ClientFormSet(request.POST, queryset=Client.objects.all())
        if formset.is_valid():
            formset.save()
            return redirect('success_url')  # Replace with your success URL
    else:
        formset = ClientFormSet(queryset=Client.objects.all())

    return render(request, 'manage_clients.html', {'formset': formset})
def manage_invoices(request):
    if request.method == 'POST':
        formset = InvoiceFormSet(request.POST)
        if formset.is_valid():
            formset.save()
            return redirect('success_url')  # Replace with your success URL
    else:
        formset = InvoiceFormSet(queryset=Invoice.objects.all())

    return render(request, 'manage_invoices.html', {'formset': formset})
def modify_and_send_file(request, invoice_id):
    try:
        invoice = Invoice.objects.get(pk=invoice_id)
        client_data = [
            {'name': 'solar_panel : ' + invoice.solar_panel.name, 'quantity': invoice.solar_panel_quantity, 'price': invoice.solar_panel_price},
            {'name': 'inverter : ' + invoice.inverter.name, 'quantity': invoice.inverter_quantity, 'price': invoice.inverter_price},
            {'name': 'structure : ' + invoice.structure.name, 'quantity': invoice.structure_quantity, 'price': invoice.structure_price},
            {'name': 'cabling : ' + invoice.cabling.name, 'quantity': invoice.cabling_quantity, 'price': invoice.cabling_price},
            {'name': 'net_metering : ' + invoice.net_metering.name, 'quantity': invoice.net_metering_quantity, 'price': invoice.net_metering_price}
        ]

        print('Invoice Details:', client_data)

        file_path = os.path.join(settings.MEDIA_ROOT, 'invoices', 'template.docx')
        if not os.path.exists(file_path):
            raise Http404("File not found.")
    except Invoice.DoesNotExist:
        raise Http404("Invoice not found.")
    def handle_table0(table):
        # Set Status
        
        row = table.rows[0].cells
        cell = row[0]
        # Set text properties
        paragraph = cell.paragraphs[0]
        paragraph.clear()  # Clear existing text in the paragraph
        run = paragraph.add_run(str(invoice.status))
        run.bold = True
        run.font.name = 'Roboto'

        # This part ensures the font name is set correctly
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Roboto')
        
        run.font.size = Pt(24)  # Adjust the size as needed
        run.font.color.rgb = RGBColor(127, 127, 127)  # Dark blue color
        # Set date
        row = table.rows[1].cells
        cell = row[4]
        # Set text properties
        paragraph = cell.paragraphs[0]
        paragraph.clear()  # Clear existing text in the paragraph
        run = paragraph.add_run(str(invoice.updated_at.strftime('%Y-%m-%d')))
        run.bold = True
        run.font.name = 'Roboto'

        # This part ensures the font name is set correctly
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Roboto')
        
        run.font.size = Pt(12)  # Adjust the size as needed
        run.font.color.rgb = RGBColor(0, 51, 153)  # Dark blue color
        # set invoice No
        row = table.rows[3].cells
        cell = row[4]
        # Set text properties
        paragraph = cell.paragraphs[0]
        paragraph.clear()  # Clear existing text in the paragraph
        run = paragraph.add_run('#' + str(invoice.id))
        run.bold = True
        run.font.name = 'Roboto'

        # This part ensures the font name is set correctly
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Roboto')
        
        run.font.size = Pt(12)  # Adjust the size as needed
        run.font.color.rgb = RGBColor(0, 51, 153)  # Dark blue color
    
    def handle_table1(table):
        row = table.rows[2].cells
        row[0].text = "Name : "+ str(invoice.name.name)
        
        # Company Name
        row = table.rows[3].cells
        row[0].text = "Company : "+ str(invoice.name.company)
        # Address
        row = table.rows[4].cells
        row[0].text = "Address : "+ str(invoice.name.city)
        # Phone
        row = table.rows[5].cells
        row[0].text = "Phone : "+ str(invoice.name.contact_number)
    
    def handle_table2(table):
        total = 0
        row_index = 1  # Assuming the first row is the header row

        for item in client_data:
            if row_index < len(table.rows):
                row = table.rows[row_index].cells
                row[1].text = str(item.get('name'))
                row[2].text = str(item.get('quantity'))
                row[3].text = str(item.get('price'))
                price = int(item.get('quantity', 0)) * int(item.get('price', 0))
                row[4].text = str(price)
                total += price
                row_index += 1
            else:
                # If there are not enough rows in the table, you can decide to add new rows if necessary
                row = table.rows[row_index].cells
                row[1].text = str(item.get('name'))
                row[2].text = str(item.get('quantity'))
                row[3].text = str(item.get('price'))
                price = int(item.get('quantity', 0)) * float(item.get('price', 0))
                row[4].text = str(price)
                total += price
            # Add Subtotal
            row = table.rows[6].cells
            row[4].text = str(total)
            # Add Discount
            row = table.rows[7].cells
            row[4].text = str(invoice.discount)
            # Add SHIPPING
            row = table.rows[10].cells
            row[4].text = str(invoice.shipping_charges)
            # Add Total
            row = table.rows[11].cells
            row[4].text = str(int(total) - int(invoice.discount)+int(invoice.shipping_charges))

    doc = Document(file_path)
    # Adding Client data
    table = doc.tables[0]
    handle_table0(table)
    
    # Adding Client data
    table = doc.tables[1]
    handle_table1(table)

    table = doc.tables[2]
    handle_table2(table)
    tmp = tempfile.NamedTemporaryFile(delete=False)
    doc.save(tmp.name)
    tmp.close()  # Manually close the file to ensure all data is written

    tmp = open(tmp.name, 'rb')
    response = FileResponse(tmp, as_attachment=True, filename=f'modified_invoice_{invoice.id}.docx')
    response.headers['Content-Disposition'] = f'attachment; filename="modified_invoice_{invoice.id}.docx"'
    return response
class GenericModelListView(ListView):
    template_name = 'generic_list.html'
    action = 'list'
    paginate_by = 5
    models = ['Client', 'SolarPanel', 'Inverter', 'Structure', 'Cabling', 'NetMetering', 'Invoice']
    
    def get_queryset(self):
        self.model = apps.get_model('crud', self.kwargs['model_name'])
        queryset = self.model.objects.all()
        queryset = self.apply_filters(queryset)
        queryset = self.apply_sort(queryset)
        return queryset
    
    def apply_filters(self, queryset):
        search_query = self.request.GET.get('q')
        date_range = self.request.GET.get('date_range')
        if search_query:
            query = Q()
            for field in self.get_searchable_fields():
                query |= Q(**{f"{field}__icontains": search_query})
            queryset = queryset.filter(query)
        if date_range:
            try:
                start_date, end_date = date_range.split(' to ')
                start_date = datetime.strptime(start_date, '%Y-%m-%d')
                end_date = datetime.strptime(end_date, '%Y-%m-%d')
                queryset = queryset.filter(created_at__range=(start_date, end_date))
            except ValueError:
                pass
        return queryset

    def get_searchable_fields(self):
        fields = []
        for field in self.model._meta.fields:
            if isinstance(field, (models.CharField, models.TextField)):
                fields.append(field.name)
            elif isinstance(field, models.ForeignKey):
                # Adjust the search to use 'name' or 'brand' for ForeignKey fields
                if hasattr(field.related_model, 'name'):
                    fields.append(f"{field.name}__name")
                elif hasattr(field.related_model, 'brand'):
                    fields.append(f"{field.name}__brand")
        return fields

    def apply_sort(self, queryset):
        sort_by = self.request.GET.get('sort_by', None)
        if sort_by:
            if sort_by.startswith('-'):
                sort_by = sort_by[1:]
            else:
                sort_by = f'-{sort_by}'
            queryset = queryset.order_by(sort_by)
        return queryset
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['model_name'] = self.kwargs['model_name']
        context['field_labels'] = self.model.get_field_labels()
        return context

class GenericModelCreateView(CreateView):

    action = 'ADD'
    def get_template_names(self):
        model_name = self.kwargs.get('model_name')
        if model_name == 'Invoice':
            return ['invoice_form.html']  # Template name for Invoice model
        else:
            return ['generic_form.html']  # Template name for other models
    

    def get_form_class(self):
        model_name = self.kwargs['model_name']
        form_class_name = self.kwargs['model_name'] + 'Form'
        print('form class name is --> ',form_class_name)
        return getattr(forms, form_class_name)
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self.request.method in ('POST', 'PUT'):
            kwargs.update({'files': self.request.FILES})
        return kwargs
    def get_context_data(self, **kwargs):
        context = super(GenericModelCreateView, self).get_context_data(**kwargs)
        context['model_name'] = self.kwargs.get('model_name', 'InvoiceModel')
        context['action'] = self.action
        return context
    def get_success_url(self):
        return reverse_lazy('generic_list', kwargs={'model_name': self.kwargs['model_name']})

class GenericModelUpdateView(UpdateView):
    
    action = 'update'
    def get_template_names(self):
        model_name = self.kwargs.get('model_name')
        if model_name == 'Invoice':
            return ['invoice_form.html']  # Template name for Invoice model
        else:
            return ['generic_form.html']  # Template name for other models

    def get_object(self, queryset=None):
        model = apps.get_model('crud', self.kwargs['model_name'])
        return model.objects.get(pk=self.kwargs['pk'])

    def get_form_class(self):
        form_class_name = self.kwargs['model_name'] + 'Form'
        return getattr(forms, form_class_name)
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self.request.method in ('POST', 'PUT'):
            kwargs.update({'files': self.request.FILES})
        return kwargs
    def get_context_data(self, **kwargs):
        context = super(GenericModelUpdateView, self).get_context_data(**kwargs)
        context['model_name'] = self.kwargs.get('model_name', 'InvoiceModel')
        context['action'] = self.action
        return context
    def get_success_url(self):  
        return reverse_lazy('generic_list', kwargs={'model_name': self.kwargs['model_name']})

class GenericModelDeleteView(DeleteView):
    template_name = 'generic_confirm_delete.html'
    action = 'delete'
    def get_context_data(self, **kwargs):
        context = super(GenericModelDeleteView, self).get_context_data(**kwargs)
        context['model_name'] = self.kwargs.get('model_name', 'InvoiceModel')
        context['action'] = self.action
        return context
    def get_object(self, queryset=None):
        model = apps.get_model('crud', self.kwargs['model_name'])
        return model.objects.get(pk=self.kwargs['pk'])

    def get_success_url(self):
        return reverse_lazy('generic_list', kwargs={'model_name': self.kwargs['model_name']})

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['model_name'] = self.kwargs['model_name']
        return context

    """
    A ViewSet that determines the queryset and serializer class dynamically
    based on the 'model_name' URL keyword argument.
    """

    def get_queryset(self):
        model_name = self.kwargs.get('model_name')
        if model_name:
            model = apps.get_model('crud', model_name)
            return model.objects.all()
        else:
            raise ImproperlyConfigured("No model specified!")

    def get_serializer_class(self):
        model_name = self.kwargs.get('model_name')
        if model_name:
            class Meta:
                model = apps.get_model('crud', model_name)
                model = model
                fields = '__all__'

            serializer_class = type(
                f"{model_name}Serializer",
                (serializers.ModelSerializer,),
                {'Meta': Meta}
            )
            return serializer_class
        else:
            raise ImproperlyConfigured("No model specified!")